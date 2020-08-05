# -*- coding: utf-8 -*-
"""
Google翻訳プログラム text_translate.py
M. Imamura 2019/7/12

"""
import tkinter as tk
import tkinter.scrolledtext as S
from tkinter import messagebox as tkMessageBox
from tkinter import filedialog as tkFileDialog
from googletrans import Translator
from docx import Document
import os
import datetime

# proxyを環境変数に設定する。
# os.environ["http_proxy"] = "10.1.10.133:8080"
# os.environ["https_proxy"] = "10.1.10.133:8080"

# ボタン1がクリックされた時の処理
def ButtonClick1():
    lines = inputbox.get('1.0', 'end -1c')      # 入力欄に入力された文字列を取得

    # 改行前の「.」、「."」、「.”」、「：」をダミー文字に置き換え
    lines = lines.replace(".\n", "XXX1").replace(".\"\n", "XXX2").replace(".”\n", "XXX3").replace(":\n", "XXX4")

    lines = lines.replace("-\n", "")            # 改行前の「-」を削除

    lines = lines.replace("\n", " ")            # 改行記号を削除

    # ダミー文字を元に戻してさらに空行を追加
    lines = lines.replace("XXX1", ".\n\n").replace("XXX2", ".\"\n\n").replace("XXX3", ".”\n\n").replace("XXX4", ":\n\n")

    lines = lines.encode('utf-8', "ignore").decode('utf-8')	# Pythonで取り扱えない文字を削除する。

    resultbox.delete('1.0', 'end')              # 整形結果欄をクリア

    resultbox.insert('1.0', lines)              # 整形結果欄に整形結果を出力

# ボタン2がクリックされた時の処理
def ButtonClick2():
    lines = resultbox.get('1.0', 'end -1c')     # 整形結果欄に入力された文字列を取得

    # 文字数を取得して画面に表示
    label = tk.Label(root, text = "文字数 "  + str(len(lines)) + "   ", font = ("Helvetica",14))
    label.place(relx = 0.55, y = 20)

    translate_box.delete('1.0', 'end')          #　翻訳結果欄をクリア

    while True:
        if len(lines) >= 5000:                  # 翻訳文字数が5000字以上の場合
            lines1 = lines[:5000].rsplit('\n\n', 1)[0]                  # 5000字以内の段落
            lines2 = lines[:5000].rsplit('\n\n', 1)[1] + lines[5000:]   # 残りの段落
    
            translator = Translator()
            lines1 = translator.translate(lines1, dest='ja').text       # Google翻訳
            translate_box.insert('end', lines1 + '\n\n')                # 翻訳結果欄に表示
    
            lines = lines2                      # 残りの段落を設定
        
        else:                                   # 翻訳文字数が5000字未満の場合
            translator = Translator()
            lines = translator.translate(lines, dest='ja').text         # Google翻訳
            translate_box.insert('end', lines)                          # 翻訳結果欄に表示

            break

# ボタン3がクリックされた時の処理
def ButtonClick3():
    edit_text = resultbox.get('1.0', 'end -1c')                         # 整形結果欄に入力された文字列を取得

    translate_text = translate_box.get('1.0', 'end -1c')                # 翻訳結果欄に入力された文字列を取得

    fTyp=[('wordファイル',"*.docx")]                                        # Word対比表テンプレートを選択 
    iDir='.'
    filename=tkFileDialog.askopenfilename(filetypes=fTyp,initialdir=iDir)
    document = Document(filename)

    for paragraph in document.paragraphs:
        paragraph.text = paragraph.text.replace("原文をここに記載する。",edit_text)        # Word対比表に原文を記載
        paragraph.text = paragraph.text.replace("訳文をここに記載する。",translate_text)   # Word対比表に訳文を記載

    paragraphs = (paragraph
                  for table in document.tables
                  for row in table.rows
                  for cell in row.cells
                  for paragraph in cell.paragraphs)
    
    for paragraph in paragraphs:
        paragraph.text = paragraph.text.replace("原文をここに記載する。",edit_text)        # Word対比表に原文を記載
        paragraph.text = paragraph.text.replace("訳文をここに記載する。",translate_text)   # Word対比表に訳文を記載

    dt_now = datetime.datetime.now()                                    # Word対比表の保存
    dt_str = str(dt_now.hour).zfill(2)+str(dt_now.minute).zfill(2)+str(dt_now.second).zfill(2)
    savefilename = filename.replace(u".docx","_replace" + dt_str + ".docx")
    document.save(savefilename)

    tkMessageBox.showinfo("作成完了",os.path.basename(savefilename) + " で保存しました。")   # 保存結果を表示
    
# メインのプログラム
root = tk.Tk()
root.geometry("1600x800")
root.title("Google翻訳プログラム")

# ラベルの設定
label1 = tk.Label(root, text = "テキストを入力", font = ("Helvetica",14))
label1.place(x = 20, y = 20)

label2 = tk.Label(root, text = "整形結果", font = ("Helvetica",14))
label2.place(relx = 0.34, y = 20)

label3 = tk.Label(root, text = "翻訳結果", font = ("Helvetica",14))
label3.place(relx = 0.67, y = 20)

# ボタンの設定
button1 = tk.Button(root, text = "整形", font = ("Helvetica",14), command = ButtonClick1)
button1.place(x = 200, y = 15)

button2 = tk.Button(root, text = "Google翻訳", font = ("Helvetica",14), command = ButtonClick2)
button2.place(relx = 0.42, y = 15)

button3 = tk.Button(root, text = "Word対比表作成", font = ("Helvetica",14), command = ButtonClick3)
button3.place(relx = 0.75, y = 15)

# 入力ボックスの設定
inputbox = S.ScrolledText(root, font = ("Helvetica",12))
inputbox.place(relheight = 0.89, relwidth = 0.32, relx = 0.01, y = 60)

# 整形結果ボックスの設定
resultbox = S.ScrolledText(root, font = ("Helvetica",12))
resultbox.place(relheight = 0.89, relwidth = 0.32, relx = 0.34, y = 60)

# 翻訳ボックスの設定
translate_box = S.ScrolledText(root, font = ("Helvetica",12))
translate_box.place(relheight = 0.89, relwidth = 0.32, relx = 0.67, y = 60)

root.mainloop()

